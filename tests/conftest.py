import pytest
from docx import Document


@pytest.fixture
def simple_docx_path(tmp_path):
    doc = Document()
    doc.add_paragraph("Hello")
    doc.add_paragraph("World")
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def empty_docx_path(tmp_path):
    doc = Document()
    path = tmp_path / "empty.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def nonexistent_docx_path():
    return "/nonexistent/path/file.docx"


@pytest.fixture
def not_docx_path(tmp_path):
    path = tmp_path / "file.txt"
    path.write_text("hello")
    return str(path)


@pytest.fixture
def rich_docx_path(tmp_path):
    doc = Document()
    doc.add_paragraph("Hello", style="Heading 1")
    doc.add_paragraph("World")
    doc.add_table(rows=2, cols=3)
    path = tmp_path / "rich.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def table_docx_path(tmp_path):
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(0, 2).text = "C"
    table.cell(1, 0).text = "D"
    table.cell(1, 1).text = "E"
    table.cell(1, 2).text = "F"
    path = tmp_path / "table.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def new_docx_path(tmp_path):
    return str(tmp_path / "new.docx")


@pytest.fixture
def new_not_docx_path(tmp_path):
    return str(tmp_path / "new.txt")
