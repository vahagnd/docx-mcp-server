"""Adding content tools for MCP server"""

from pathlib import Path
from typing import Optional

from docx.shared import Inches, Pt

from src.utils import load, resolve_align, save


def add_paragraph(
    path: str,
    text: str,
    style: str = "Normal",
    bold: bool = False,
    italic: bool = False,
    font_size: Optional[int] = None,
    alignment: Optional[str] = None,
) -> str:
    """
    Append a paragraph to an existing .docx file.

    Args:
        path:      Path to the .docx file.
        text:      Paragraph text content.
        style:     Paragraph style name (default "Normal").
        bold:      Apply bold to the run (default False).
        italic:    Apply italic to the run (default False).
        font_size: Font size in points (e.g. 12).
        alignment: "left" | "center" | "right" | "justify".

    Returns:
        Confirmation with the index of the new paragraph.
    """
    doc = load(path)
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if alignment:
        para.alignment = resolve_align(alignment)
    save(doc, path)
    return f"Paragraph added at index {len(doc.paragraphs) - 1}."


def add_heading(path: str, text: str, level: int = 1) -> str:
    """
    Append a heading paragraph to an existing .docx file.

    Args:
        path:  Path to the .docx file.
        text:  Heading text.
        level: Heading level 1–6 (default 1).

    Returns:
        Confirmation message.
    """
    if not 1 <= level <= 6:
        raise ValueError(f"Level must be 1–6, got {level}.")
    doc = load(path)
    doc.add_heading(text, level=level)
    save(doc, path)
    return f'Heading level {level} added: "{text}".'


def add_table(
    path: str,
    data: list[list[str]],
    has_header_row: bool = True,
    style: str = "Table Grid",
) -> str:
    """
    Append a table to an existing .docx file.

    Args:
        path:           Path to the .docx file.
        data:           2-D list of strings; each inner list is a row.
                        Example: [["Name", "Age"], ["Alice", "30"]]
        has_header_row: Bold the first row (default True).
        style:          Word table style name (default "Table Grid").

    Returns:
        Confirmation with table dimensions.
    """
    if not data or not data[0]:
        raise ValueError("data must be a non-empty 2-D list.")
    doc = load(path)
    rows, cols = len(data), max(len(r) for r in data)
    tbl = doc.add_table(rows=rows, cols=cols)
    try:
        tbl.style = style
    except KeyError:
        pass
    for r_idx, row_data in enumerate(data):
        for c_idx, cell_text in enumerate(row_data):
            cell = tbl.cell(r_idx, c_idx)
            cell.text = cell_text
            if has_header_row and r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    save(doc, path)
    return f"Table ({rows}×{cols}) added."


def add_picture(
    path: str,
    image_path: str,
    width_inches: Optional[float] = None,
) -> str:
    """
    Append an image to an existing .docx file.

    Args:
        path:         Path to the .docx file.
        image_path:   Path to the image file (PNG, JPEG, GIF, BMP, TIFF).
        width_inches: Desired width in inches. Height scales proportionally.
                      Omit to use the image's native size.

    Returns:
        Confirmation message.
    """
    if not Path(image_path).exists():
        raise FileNotFoundError(f"Image not found: {image_path}")
    doc = load(path)
    doc.add_picture(image_path, width=Inches(width_inches) if width_inches else None)
    save(doc, path)
    return f"Image '{image_path}' added."


def add_page_break(path: str) -> str:
    """
    Append a page break to an existing .docx file.

    Args:
        path: Path to the .docx file.

    Returns:
        Confirmation message.
    """
    doc = load(path)
    doc.add_page_break()
    save(doc, path)
    return "Page break added."
